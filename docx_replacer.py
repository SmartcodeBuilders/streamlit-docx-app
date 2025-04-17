# Imports Section
from docx import Document
import pandas as pd
from io import BytesIO
import streamlit as st
from docx.shared import Inches
import tempfile
import os
from PIL import Image
import io
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

from utils import get_image_from_gdrive


# List of registered doctors with their names and professional numbers
registered_doctors = [
    {
        "doctor_name": "Antonio Buzido Jimenez",
        "number": "Médico colegiado Nº 9800 de Sevilla",
        "signature_url": "https://drive.google.com/file/d/1uiMxUdsYmBYXyJwKSya7J0NAZsShmNYK/view",
        "signature_size":1
    },
    {
        "doctor_name": "Dra. Paz Marian Casal",
        "number": "Médico colegiado Nº 987 de Sevilla",
        "signature_url": "https://drive.google.com/file/d/1uiMxUdsYmBYXyJwKSya7J0NAZsShmNYK/view",
        "signature_size":2
    },
]


# Core Functions - - - -
@st.dialog("Generar Informe Médico", width="large")
def preview_file(final_df, default_open_status):
    """
    Displays a dialog for generating medical reports using a .docx template.

    This dialog allows the user to:
    - Upload a template (.docx) file.
    - Select a doctor from a list of registered doctors.
    - Choose a column from a DataFrame to be extracted into the report.
    - Provide additional details such as case number, documentation given, and documentation not given.

    Args:
        final_df (pd.DataFrame): The DataFrame containing the data that will populate the report.
    """

    # Simple toggle using session_state
    st.session_state.show_download = default_open_status

    col1, col2 = st.columns([2, 2])  # Adjust width ratio for the columns

    # File uploader for the user to select a Word template
    template_file = st.file_uploader(
        "Por favor selecciona la plantilla deseada", type="docx"
    )

    max_col = final_df.shape[1] - 1  # Maximum column index
    column_indices = list(range(0, max_col + 1))  # List of available column indices

    # Create a list of doctor names for the dropdown
    doctor_names = [doc["doctor_name"] for doc in registered_doctors]

    # Initialize session state for tracking changes
    if "doctor" not in st.session_state:
        st.session_state.doctor = ""
    if "selected_column" not in st.session_state:
        st.session_state.selected_column = ""
    if "expedient_number" not in st.session_state:
        st.session_state.expedient_number = ""
    if "documentation_given" not in st.session_state:
        st.session_state.documentation_given = ""
    if "documentation_not_given" not in st.session_state:
        st.session_state.documentation_not_given = ""

    # Selectbox for choosing the doctor
    with col1:
        doctor = st.selectbox("Doctor", doctor_names)
    # Selectbox for choosing the column to extract from the DataFrame
    with col2:
        selected_column = st.selectbox(
            "Seleccione una columna", column_indices, index=len(column_indices) - 1
        )

    (col1,) = st.columns([2])  # Adjust width ratio for the columns
    # Text inputs for additional information required in the report
    with col1:
        expedient_number = st.text_input("Número de Expediente")

    col1, col2 = st.columns([2, 2])  # Adjust width ratio for the columns
    with col1:
        documentation_given = st.text_area("Documentación Aportada", height=68)
    with col2:
        documentation_not_given = st.text_area("Documentación no Aportada", height=68)

    col1, col2 = st.columns([2, 1])  # Adjust width ratio for the columns
    
    # Check if any input has changed
    if doctor != st.session_state.doctor:
        st.session_state.show_download = False
        st.session_state.doctor = doctor

    if selected_column != st.session_state.selected_column:
        st.session_state.show_download = False
        st.session_state.selected_column = selected_column

    if expedient_number != st.session_state.expedient_number:
        st.session_state.show_download = False
        st.session_state.expedient_number = expedient_number

    if documentation_given != st.session_state.documentation_given:
        st.session_state.show_download = False
        st.session_state.documentation_given = documentation_given

    if documentation_not_given != st.session_state.documentation_not_given:
        st.session_state.show_download = False
        st.session_state.documentation_not_given = documentation_not_given

    # Find the selected doctor in the list of registered doctors
    matched_doctor = next(
        (doc for doc in registered_doctors if doc["doctor_name"] == doctor), "N/A"
    )

    # Collect the additional information to be used in the document
    additional_documentation = {
        "{{Doctor}}": doctor,
        "{{Numero de colegiado}}": matched_doctor["number"],
        "{{Doctor Identification}}": matched_doctor["number"],
        "{{Expediente}}": expedient_number,
        "{{Documentación aportada}}": documentation_given,
        "{{Documentación no aportada}}": documentation_not_given,
    }

    if st.button("Procesar información", use_container_width=True):
        st.session_state.show_download = True

    # Trigger document generation and offer download options if the template is uploaded
    if st.session_state.show_download == True:
        if template_file:
            fill_and_offer_multiple_downloads(
                final_df, selected_column, template_file, additional_documentation, matched_doctor["signature_url"], matched_doctor["signature_size"]
            )


def fill_and_offer_multiple_downloads(
    df: pd.DataFrame, column_index: int, template_path: str, extra_information: dict, signature_url, signature_size
):
    """
    Fills a Word template with data from a given DataFrame column and shows download buttons for both DOCX and PDF formats.

    This function:
    - Extracts the data from the DataFrame column.
    - Replaces placeholders in the Word document with the extracted data and additional information.
    - Provides download buttons for both DOCX and PDF formats.

    Args:
        df (pd.DataFrame): The DataFrame containing data to be inserted into the report.
        column_index (int): The index of the column in the DataFrame to extract.
        template_path (str): Path to the .docx template that will be used to generate the report.
        extra_information (dict): A dictionary of additional information to be inserted into the template.
    """

    # Layout: two columns for the download buttons
    col1, col2 = st.columns(2)

    # Extract the column data to be used in the template
    column_data = df[column_index]

    # Prepare the replacements for placeholders in the template
    replacements = {
        f"{{{{{key}}}}}": (
            (str(value).split(" Hora: ")[0] if key == "Fecha siniestro" else str(value))
            if pd.notnull(value)
            else ""
        )
        for key, value in column_data.items()
    }

    # Special handling for the 'Fecha siniestro' field to separate date and time
    if "Fecha siniestro" in column_data and pd.notnull(column_data["Fecha siniestro"]):
        fecha_value = column_data["Fecha siniestro"]
        if " Hora: " in fecha_value:
            replacements["{{Fecha Siniestro}}"] = fecha_value.split(" Hora: ")[0]
            replacements["{{Hora}}"] = fecha_value.split(" Hora: ")[1]

    if "Numero de documento" in column_data and pd.notnull(
        column_data["Numero de documento"]
    ):
        file_name = column_data["Numero de documento"]

    if "Nombre y apellidos" in column_data and pd.notnull(
        column_data["Nombre y apellidos"]
    ):
        file_name = f"{file_name} {column_data['Nombre y apellidos']}"

    if "{{Expediente}}" in extra_information:
        file_name = f"{file_name} {extra_information['{{Expediente}}']}"

    # Load the template Word document
    doc = Document(template_path)
    
    
    for para in doc.paragraphs:
        for key, val in replacements.items():
            if key in para.text:
                para.text = para.text.replace(key, val)
                
    for para in doc.paragraphs:
        for key, val in extra_information.items():
            if key in para.text:
                para.text = para.text.replace(key, val)
    

    
    # Replace placeholders in the tables of the document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for key, val in replacements.items():
                            if key in run.text and isinstance(val, str):
                                run.text = run.text.replace(key, val)
                        for key, val in extra_information.items():
                            if key in run.text and isinstance(val, str):
                                run.text = run.text.replace(key, val)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, val)
                for key, val in extra_information.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, val)
    
    # Replace images
    for i, para in enumerate(doc.paragraphs):
        if "{{signature image}}" in para.text:
            try:
                # 1. Eliminar el párrafo con el placeholder
                p = para._element
                parent = p.getparent()
                index = parent.index(p)
                parent.remove(p)

                # 2. Crear nuevo párrafo con la imagen alineada a la derecha
                new_paragraph = doc.add_paragraph()
                new_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                signature_bytes = get_image_from_gdrive(signature_url)
                signature_bytes.seek(0)
                run = new_paragraph.add_run()
                run.add_picture(signature_bytes, width=Inches(signature_size))

                # 3. Insertar el nuevo párrafo en la posición original
                parent.insert(index, new_paragraph._element)
            except (AttributeError, ValueError, FileNotFoundError) as e:
                st.warning("⚠️ No se pudo cargar la imagen de la firma. Asegúrate de que el enlace sea válido.")
        
            break
    
    # Save the modified document to an in-memory buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Download button for DOCX format
    with col1:
        st.download_button(
            label=f"📄 Descargar Informe en formato .docx",
            data=buffer,
            file_name=f"{file_name}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
