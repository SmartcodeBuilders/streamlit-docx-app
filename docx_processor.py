from typing import IO
import re
import pandas as pd
from docx import Document

class WordProcessor:
    def __init__(self, doc_data: IO[bytes]):
        
        
        self.tables = []
        self.full_text = []
        self.doc = Document(doc_data)

        # DataFrame for personal details and initial medical information (tables 1-5)
        self.df = pd.DataFrame()
        # DataFrame for first medical visit, secuelas, and lawyer information (tables 6, 7, and 9)
        self.first_medical_visit = pd.DataFrame()
        self.next_medical_visits = []
        self._process_tables()
        self.populate_first_table_dataframe()
        self.populate_second_table_dataframe()
        self.populate_third_table_dataframe()
        self.populate_fourth_table_dataframe()
        self.populate_fifth_table_dataframe()
        self.populate_first_medical_visit_dataframe()
        self.populate_next_medical_visits_dataframe()
        self.populate_visits_from_full_text()
        
    
    def _process_tables(self) -> None:
        for table in self.doc.tables:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                table_text.append(row_text)
            self.tables.append(table_text)
            
            
        all_paragraphs = []
        all_paragraphs.extend(self.doc.paragraphs)
        
        for paragraph in all_paragraphs:
            if paragraph.text.strip():
                self.full_text.append(paragraph.text)
    
    def populate_first_table_dataframe(self) -> None:
        """
        Processes the first table and populates self.df with company-related fields:
        "Compañía", "Fecha siniestro", "Hora", "Lugar de la visita", "Fecha visita", "Nombre del Doctor"
        """
        datos_compañia = [
            "Compañía",
            "Fecha siniestro",
            "Hora",
            "Lugar de la visita",
            "Fecha visita",
            "Nombre del Doctor"
        ]
        data = {field: None for field in datos_compañia}
        
        if not self.tables:
            self.df = pd.DataFrame()
            return
        
        for row in self.tables[0]:
            for cell_text in row:
                segments = re.split(r'\s{2,}', cell_text)
                for seg in segments:
                    if ':' in seg:
                        key, val = seg.split(':', 1)
                        key = key.strip()
                        val = val.strip()
                        if key == "Fecha visita":
                            match = re.match(r'(\d{2}/\d{2}/\d{2,4})\s*(.*)', val)
                            if match:
                                fecha_visita, doctor = match.groups()
                                data["Fecha visita"] = fecha_visita.strip()
                                data["Nombre del Doctor"] = doctor.strip() if doctor else None
                            else:
                                data["Fecha visita"] = val
                        else:
                            if key in data:
                                data[key] = val
        self.df = pd.DataFrame([data])
    
    def populate_second_table_dataframe(self) -> None:
        """
        Processes the second table and updates self.df with injured person fields:
        "Nombre y apellidos", "Condición", "Domicilio", "NIF", "Población", "Teléfono (FyM)",
        "C.P.", "Edad", "Fecha nacimiento", "Provincia", "Sexo", "Lateralidad", "Profesión",
        "Nivel s.e.", "Puesto de trabajo / ocupación", "Deportes", "Federado",
        "Situación laboral en el momento del accidente", "Actividades de ocio", "Mail",
        "Protección", "¿Agravación por no uso protección?"
        """
        datos_del_lesionado = [
            "Nombre y apellidos",
            "Condición",
            "Domicilio",
            "NIF",
            "Población",
            "Teléfono (FyM)",
            "C.P.",
            "Edad",
            "Fecha nacimiento",
            "Provincia",
            "Sexo",
            "Lateralidad",
            "Profesión",
            "Nivel s.e.",
            "Puesto de trabajo / ocupación",
            "Deportes",
            "Federado",
            "Situación laboral en el momento del accidente",
            "Actividades de ocio",
            "Mail",
            "Protección",
            "¿Agravación por no uso protección?",
        ]
        data = {field: None for field in datos_del_lesionado}
        
        if len(self.tables) < 2:
            for field in data:
                if not data[field]:
                    data[field] = "-"
            self.df = self.df.assign(**data)
            return
        
        for row in self.tables[1]:
            for cell_text in row:
                cleaned = cell_text.replace('\n', ' ')
                cleaned = re.sub(r'[\u2002\u2003\u2004\u2005\u2006\u2007\u2008\u2009\u200a]+', ' ', cleaned)
                segments = re.split(r'\s{2,}', cleaned)
                for seg in segments:
                    if ':' in seg:
                        key, val = seg.split(':', 1)
                        key = key.strip()
                        val = val.strip().replace(" ¿Agravación por no uso protección?:", "")
                        if key in data:
                            data[key] = val
        
        for key in data:
            if not data[key]:
                data[key] = "-"
        
        self.df = self.df.assign(**data)
    
    def populate_third_table_dataframe(self) -> None:
        """
        Processes the third table and updates self.df with familial fields:
        "Estado civil", "Nº de Hijos", "Menores", "Miembros unidad familiar", "Miembros discapacitados"
        """
        datos_familiares = [
            "Estado civil",
            "Nº de Hijos",
            "Menores",
            "Miembros unidad familiar",
            ">18 años",
            "<18 años",
            "Miembros discapacitados"
        ]
        data = {field: None for field in datos_familiares}
        
        if len(self.tables) < 3:
            for field in data:
                if not data[field]:
                    data[field] = "-"
            self.df = self.df.assign(**data)
            return
        
        for row in self.tables[2]:
            for cell_text in row:
                cleaned = cell_text.replace('\n', ' ')
                cleaned = re.sub(r'[\u2002\u2003\u2004\u2005\u2006\u2007\u2008\u2009\u200a]+', ' ', cleaned)
                segments = re.split(r'\s{2,}', cleaned)
                for seg in segments:
                    if ':' in seg:
                        key, val = seg.split(':', 1)
                        key = key.strip()
                        val = val.strip()
                        if key in data:
                            data[key] = val
        
        for key in data:
            if not data[key]:
                data[key] = "-"
        
        self.df = self.df.assign(**data)
    
    def populate_fourth_table_dataframe(self) -> None:
        """
        Processes the fourth table and updates self.df with:
        Fields: ['Tipo', 'Fecha ingreso', 'Fecha alta', 'Nº Historial Clínico']
        Assumes the table has two rows: first row as headers, second row as values.
        """
        expected_fields = ['Tipo', 'Fecha ingreso', 'Fecha alta', 'Nº Historial Clínico']
        data = {field: "-" for field in expected_fields}
        
        if len(self.tables) < 4:
            self.df = self.df.assign(**data)
            return
        
        table4 = self.tables[3]
        if len(table4) >= 2:
            headers = table4[0]
            values = table4[1]
            for header, value in zip(headers, values):
                header_clean = header.strip()
                value_clean = value.strip()
                if header_clean in expected_fields:
                    data[header_clean] = value_clean if value_clean else "-"
        
        self.df = self.df.assign(**data)
    
    def populate_fifth_table_dataframe(self) -> None:
        """
        Processes the fifth table and updates self.df with:
        Fields: ['Códigos', 'Diagnóstico']
        Assumes the table has two rows: first row as headers, second row as values.
        """
        expected_fields = ['Códigos', 'Diagnóstico']
        data = {field: "-" for field in expected_fields}
        
        if len(self.tables) < 5:
            self.df = self.df.assign(**data)
            return
        
        table5 = self.tables[4]
        if len(table5) >= 2:
            headers = table5[0]
            values = table5[1]
            for header, value in zip(headers, values):
                header_clean = header.strip()
                value_clean = value.strip()
                if header_clean in expected_fields:
                    data[header_clean] = value_clean if value_clean else "-"
        
        self.df = self.df.assign(**data)
    
    def populate_first_medical_visit_dataframe(self) -> None:
        """
        Processes table 6, table 7, and table 9, merging their information into a single DataFrame 
        (self.first_medical_visit) with the following fields:
        
        From table 6 (first medical visit):
            - "Lesiones muy graves" (from "Muy graves")
            - "Lesiones graves" (from "Graves")
            - "Lesiones moderados" (from "Moderados")
            - "Lesiones basicos" (from "Básicos")
            - "Fecha alta" (from "Fecha alta")
            - "Motivos variacion fecha final" (from "Motivos variación de fecha inicial")
        
        From table 7 (secuelas):
            - "Codigo Secuela" (from "Código")
            - "Descripción secuela" (from "Descripción secuela")
            - "analogía secuela" (from "Analogía")
            - "rango secuela" (from "Rango")
            - "prev/defin secuela" (from "Prev./Defin.")
            - "puntuación secuela" (from "Puntuación")
        
        From table 9 (lawyer information):
            - "Nombre abogado" (from "Nombre abogado")
            - "Telefono abogado" (from "Teléfono")
        """
        # --- Process table 6 ---
        mapping6 = {
            "Muy graves": "Lesiones muy graves",
            "Graves": "Lesiones graves",
            "Moderados": "Lesiones moderados",
            "Básicos": "Lesiones basicos",
            "Fecha alta": "Fecha alta",
            "Motivos variación de fecha inicial": "Motivos variacion fecha final"
        }
        data6 = {new_key: "-" for new_key in mapping6.values()}
        if len(self.tables) >= 6:
            table6 = self.tables[5]
            for row in table6:
                for i, cell_text in enumerate(row):
                    cleaned = cell_text.replace('\n', ' ').strip()
                    cleaned = re.sub(r'[\u2002\u2003\u2004\u2005\u2006\u2007\u2008\u2009\u200a]+', ' ', cleaned)
                    if ':' in cleaned:
                        key_part, val_part = cleaned.split(':', 1)
                        key_part = key_part.strip()
                        val_part = val_part.strip()
                        if key_part in mapping6:
                            if not val_part and i+1 < len(row):
                                next_cell = row[i+1].replace('\n', ' ').strip()
                                next_cell = re.sub(r'[\u2002\u2003\u2004\u2005\u2006\u2007\u2008\u2009\u200a]+', ' ', next_cell)
                                if next_cell and ':' not in next_cell:
                                    val_part = next_cell
                            data6[mapping6[key_part]] = val_part if val_part else "-"
                    else:
                        if data6["Fecha alta"] == "-" and i > 0 and "Fecha alta" in row[i-1]:
                            data6["Fecha alta"] = cleaned if cleaned else "-"
        
        # --- Process table 7 ---
        mapping7 = {
            "Código": "Codigo Secuela",
            "Descripción secuela": "Descripción secuela",
            "Analogía": "analogía secuela",
            "Rango": "rango secuela",
            "Prev./Defin.": "prev/defin secuela",
            "Puntuación": "puntuación secuela"
        }
        data7 = {new_key: "-" for new_key in mapping7.values()}
        if len(self.tables) >= 7:
            table7 = self.tables[6]
            if len(table7) >= 2:
                headers = table7[0]
                values = table7[1]
                for header, value in zip(headers, values):
                    header_clean = header.strip()
                    value_clean = value.replace('\n', ' ')
                    value_clean = re.sub(r'[\u2002\u2003\u2004\u2005\u2006\u2007\u2008\u2009\u200a]+', ' ', value_clean).strip()
                    if header_clean in mapping7:
                        data7[mapping7[header_clean]] = value_clean if value_clean else "-"
                        
                        
        # --- Process table 8 ---
        data8 = {"Perdida c vida: Grado y razonarlo": "-"}
        if len(self.tables) >= 8:
            table8 = self.tables[7]
            for row in table8:
                for cell_text in row:
                    cleaned = cell_text.replace('\n', ' ')
                    cleaned = re.sub(r'[\u2002\u2003\u2004\u2005\u2006\u2007\u2008\u2009\u200a]+', ' ', cleaned)
                    if "Grado y razonarlo:" in cleaned or "Notas:" in cleaned:
                        _, val = cleaned.split(":", 1)
                        data8["Perdida c vida: Grado y razonarlo"] = val.strip() if val.strip() else "-"
        
        
        # --- Process table 9 ---
        mapping9 = {
            "Nombre abogado": "Nombre abogado",
            "Teléfono": "Telefono abogado"
        }
        data9 = {new_key: "-" for new_key in mapping9.values()}
        if len(self.tables) > 9:
            table9 = self.tables[9]
            for row in table9:
                for cell_text in row:
                    cleaned = cell_text.replace('\n', ' ')
                    cleaned = re.sub(r'[\u2002\u2003\u2004\u2005\u2006\u2007\u2008\u2009\u200a]+', ' ', cleaned)
                    if ':' in cleaned:
                        key, val = cleaned.split(':', 1)
                        key = key.strip()
                        val = val.strip()
                        if key in mapping9:
                            data9[mapping9[key]] = val if val else "-"
        
        # --- Merge the dictionaries ---
        merged_data = {}
        merged_data.update(data6)
        merged_data.update(data7)
        merged_data.update(data8)
        merged_data.update(data9)
        
        self.first_medical_visit = pd.DataFrame([merged_data])



    def populate_next_medical_visits_dataframe(self) -> None:
        """
        Processes subsequent visit tables in pairs: [12, 13], [14, 15], etc.
        Each pair is processed similarly to the second visit function:
        
        From the first table in each pair, it extracts:
            - "Lesiones muy graves" (from "Muy graves")
            - "Lesiones graves" (from "Graves")
            - "Lesiones moderados" (from "Moderados")
            - "Lesiones basicos" (from "Básicos")
            - "Fecha alta" (from "Fecha alta")
            - "Motivos variacion fecha final" (from "Motivos variación de fecha inicial")
        
        From the second table in each pair, it extracts:
            - "Codigo Secuela" (from "Código")
            - "Descripción secuela" (from "Descripción secuela")
            - "analogía secuela" (from "Analogía")
            - "rango secuela" (from "Rango")
            - "prev/defin secuela" (from "Prev./Defin.")
            - "puntuación secuela" (from "Puntuación")
        
        The merged data for each visit is stored as a DataFrame in the list self.next_medical_visits.
        """
        self.next_medical_visits = []
        mapping6 = {
            "Muy graves": "Lesiones muy graves",
            "Graves": "Lesiones graves",
            "Moderados": "Lesiones moderados",
            "Básicos": "Lesiones basicos",
            "Fecha alta": "Fecha alta",
            "Motivos variación de fecha inicial": "Motivos variacion fecha final"
        }
        mapping7 = {
            "Código": "Codigo Secuela",
            "Descripción secuela": "Descripción secuela",
            "Analogía": "analogía secuela",
            "Rango": "rango secuela",
            "Prev./Defin.": "prev/defin secuela",
            "Puntuación": "puntuación secuela"
        }
        i = 10
        while i < len(self.tables):
            found = False
            for row in self.tables[i]:
                for cell in row:
                    if "Muy graves:" in cell:
                        found = True
                        break
                if found:
                    break
            if found:
                break
            i += 1
        while i + 1 < len(self.tables):
            # Process first table in the pair (e.g., table 12, 14, etc.)
            data_first = {v: "-" for v in mapping6.values()}
            fecha_alta_parts = []
            motivos_parts = []
            table_first = self.tables[i]
            
            for row_idx, row in enumerate(table_first):
                for cell in row:
                    cleaned = cell.replace('\n', ' ').strip()
                    cleaned = re.sub(r'[\u2002\u2003\u2004\u2005\u2006\u2007\u2008\u2009\u200a]+', ' ', cleaned)
                    if ':' in cleaned:
                        key, val = cleaned.split(':', 1)
                        key = key.strip()
                        val = val.strip()
                        if key in mapping6 and val:
                            data_first[mapping6[key]] = val
                    else:
                        # Accumulate non-colon cells for specific rows
                        if row_idx == 1:
                            fecha_alta_parts.append(cleaned)
                        elif row_idx == 2:
                            motivos_parts.append(cleaned)
            if fecha_alta_parts:
                data_first["Fecha alta"] = " ".join(fecha_alta_parts)
            if motivos_parts:
                data_first["Motivos variacion fecha final"] = " ".join(motivos_parts)
            # Process second table in the pair (e.g., table 13, 15, etc.)
            data_second = {v: "-" for v in mapping7.values()}
            table_second = self.tables[i+1]
            if len(table_second) >= 2:
                headers = table_second[0]
                values = table_second[1]
                for header, value in zip(headers, values):
                    header_clean = header.strip()
                    value_clean = value.replace('\n', ' ').strip()
                    value_clean = re.sub(r'[\u2002\u2003\u2004\u2005\u2006\u2007\u2008\u2009\u200a]+', ' ', value_clean)
                    if header_clean in mapping7 and value_clean:
                        data_second[mapping7[header_clean]] = value_clean
            merged_data = {}
            merged_data.update(data_first)
            merged_data.update(data_second)
            self.next_medical_visits.append(pd.DataFrame([merged_data]))
            i += 3


    def populate_visits_from_full_text(self) -> None:
        """
        Splits self.full_text into blocks using lines that contain "Próxima visita:".
        
        For the first block (first visit):
        - Extract "Antecedentes médicos del lesionado": the element immediately following the header.
        - Extract treatment using the first found header among 
            "Tratamiento y evolución. Exploraciones complementarias" or "Evolución"
            by joining subsequent elements until "Estado actual y exploración física" is reached.
        - For the first visit, "Fecha de consulta extra" is left empty.
        - Extract "Estado actual y exploración física" by joining subsequent elements until "Relación de causalidad" is reached.
        The resulting data is merged with any existing row in self.first_medical_visit.
        
        For each subsequent block (extra visits):
        - Do not extract antecedentes.
        - Look for a header among "Tratamiento y evolución. Exploraciones complementarias" and "Evolución".
            If "Evolución" is found, set "Fecha de consulta extra" to the first 10 characters of the element immediately preceding it.
            Otherwise, leave "Fecha de consulta extra" empty.
        - Extract treatment text from after the header until "Estado actual y exploración física" is encountered.
        - Extract "Estado actual y exploración física" by joining subsequent elements until "Relación de causalidad" is encountered.
        For each extra visit, merge the extracted data with any existing row in the corresponding DataFrame in self.next_medical_visits.
        """
        if not hasattr(self, "full_text") or not isinstance(self.full_text, list):
            return

        # Split full_text into blocks using "Próxima visita:" as delimiter
        blocks = []
        current_block = []
        for line in self.full_text:
            stripped = line.strip()
            if "Próxima visita:" in stripped:
                if current_block:
                    blocks.append(current_block)
                current_block = []
            else:
                current_block.append(stripped)
        if current_block:
            blocks.append(current_block)

        ################ Process first block (first visit) ################
        first_visit = {}
        if blocks:
            block = blocks[0]
            # Extract "Antecedentes médicos del lesionado"
            try:
                idx = block.index("Antecedentes médicos del lesionado")
                first_visit["Antecedentes médicos del lesionado"] = block[idx + 1] if idx + 1 < len(block) else "-"
            except ValueError:
                first_visit["Antecedentes médicos del lesionado"] = "-"
            ## NEW CODE ##
            # Extract "Descripción del accidente"
            try:
                accident_idx = block.index("Descripción del accidente")
                j = accident_idx + 1
                accident_parts = []
                while j < len(block) and block[j] != "Datos asistenciales":
                    accident_parts.append(block[j])
                    j += 1
                accident_text = " ".join(accident_parts)
                first_visit["Descripción del accidente"] = accident_text
            except ValueError:
                first_visit["Descripción del accidente"] = "-"
            ## END NEW CODE ##
            # Extract treatment using either header.
            treat_idx = None
            treat_header = None
            for header in ["Tratamiento y evolución. Exploraciones complementarias", "Evolución"]:
                try:
                    idx = block.index(header)
                    if treat_idx is None or idx < treat_idx:
                        treat_idx = idx
                        treat_header = header
                except ValueError:
                    continue
            if treat_idx is not None:
                j = treat_idx + 1
                treatment_parts = []
                while j < len(block) and block[j] != "Estado actual y exploración física":
                    treatment_parts.append(block[j])
                    j += 1
                treatment_text = "\n\n".join(treatment_parts)
                first_visit["Tratamiento y evolución. Exploraciones complementarias"] = treatment_text
                first_visit["Fecha de consulta extra"] = ""  # Not populated in first visit.
            else:
                first_visit["Tratamiento y evolución. Exploraciones complementarias"] = "-"
                first_visit["Fecha de consulta extra"] = ""
            # Extract state information.
            try:
                state_idx = block.index("Estado actual y exploración física")
                j = state_idx + 1
                state_parts = []
                while j < len(block) and block[j] != "Relación de causalidad":
                    state_parts.append(block[j])
                    j += 1
                state_text = "\n\n".join(state_parts)
                
                ## NEW CODE ##
                # Extract Relación de causalidad
                if j < len(block) and block[j] == "Relación de causalidad":
                    j += 1  # Move past the header
                    causalidad_parts = []
                    while j < len(block) and block[j] != "Lesiones temporales" and not "(exclusión, cronológico, topográfico, intensidad)" in block[j]:
                        causalidad_parts.append(block[j])
                        j += 1
                    
                    # Skip the exclusion text line
                    if j < len(block) and "(exclusión, cronológico, topográfico, intensidad)" in block[j]:
                        j += 1
                        # Continue collecting text after the exclusion text until Lesiones temporales
                        while j < len(block) and block[j] != "Lesiones temporales":
                            causalidad_parts.append(block[j])
                            j += 1
                    
                    first_visit["Relación de causalidad"] = " ".join(causalidad_parts)
                else:
                    first_visit["Relación de causalidad"] = "-"
                 ## END NEW CODE ##
                # Initialize with default values
                first_visit["HISTORIA ACTUAL"] = "-"
                first_visit["EXPLORACION FISICA"] = "-" 
                first_visit["Pruebas complementarias"] = "-"
                
                # Try to extract each section
                sections = state_text.split("HISTORIA ACTUAL:")
                if len(sections) > 1:
                    remaining = sections[1].strip()
                    historia = remaining.split("EXPLORACION FISICA:")[0].strip()
                    first_visit["HISTORIA ACTUAL"] = historia
                    
                    sections = remaining.split("EXPLORACION FISICA:")
                    if len(sections) > 1:
                        remaining = sections[1].strip()
                        exploracion = remaining.split("Pruebas complementarias:")[0].strip()
                        first_visit["EXPLORACION FISICA"] = exploracion
                        
                        sections = remaining.split("Pruebas complementarias:")
                        if len(sections) > 1:
                            pruebas = sections[1].strip()
                            first_visit["Pruebas complementarias"] = pruebas
            except ValueError:
                first_visit["Estado actual y exploración física"] = "-"
                
            # Extract lesiones temporales information
            try:
                lesiones_idx = block.index("Lesiones temporales")
                j = lesiones_idx + 1
                
                # Initialize with default values
                first_visit["Intervenciones quirúrgicas"] = "-"
                first_visit["Patrimonial. Daño emergente (se indemniza su importe)"] = "-"
                
                while j < len(block) and block[j] != "Secuelas. Básico":
                    text = block[j]
                    
                    # Check for "Intervenciones quirúrgicas"
                    if text.startswith("Intervenciones quirúrgicas"):
                        parts = []
                        if ":" in text:
                            # Extract value after colon in same line
                            value = text.split(":", 1)[1].strip()
                            if value:
                                parts.append(value)
                        # Also collect text from following lines
                        j += 1
                        while j < len(block) and not block[j].startswith("Patrimonial. Daño emergente"):
                            parts.append(block[j])
                            j += 1
                        first_visit["Intervenciones quirúrgicas"] = " ".join(parts) if parts else "-"
                        continue # Skip j increment since we already moved forward
                        
                    # Check for "Patrimonial. Daño emergente" 
                    elif text.startswith("Patrimonial. Daño emergente"):
                        parts = []
                        if ":" in text:
                            # Extract value after colon in same line
                            value = text.split(":", 1)[1].strip()
                            if value:
                                parts.append(value)
                        # Also collect text from following lines
                        j += 1
                        while j < len(block) and block[j] != "Secuelas. Básico":
                            parts.append(block[j])
                            j += 1
                        first_visit["Patrimonial. Daño emergente (se indemniza su importe)"] = "\n\n".join(parts) if parts else "-"
                        continue # Skip j increment since we already moved forward
                    
                    j += 1
                    
            except ValueError:
                first_visit["Intervenciones quirúrgicas"] = "-"
                first_visit["Patrimonial. Daño emergente (se indemniza su importe)"] = "-"
            
            # Extract Valoración Total Secuelas information
            try:
                j = 0
                while j < len(block):
                    if "Valoración Total Secuelas" in block[j]:
                        text = block[j]
                        if ":" in text:
                            value = text.split(":", 1)[1].strip()
                            first_visit["Valoración Total Secuelas"] = value if value else "-"
                            
                            # Check next line for "Motivos variación"
                            if j+1 < len(block) and "Motivos variación" in block[j+1]:
                                motivos_text = block[j+1]
                                if ":" in motivos_text:
                                    motivos_value = motivos_text.split(":", 1)[1].strip()
                                    first_visit["Motivos variación"] = motivos_value if motivos_value else "-"
                        else:
                            first_visit["Valoración Total Secuelas"] = "-"
                        break
                    j += 1
                    
                if "Valoración Total Secuelas" not in first_visit:
                    first_visit["Valoración Total Secuelas"] = "-"
                
                if "Motivos variación" not in first_visit:
                    first_visit["Motivos variación"] = "-"
                    
            except Exception:
                first_visit["Valoración Total Secuelas"] = "-"
                first_visit["Motivos variación"] = "-"
            
            # Extract aclaraciones information
            try:
                aclaraciones_idx = block.index("Aclaraciones:")
                j = aclaraciones_idx + 1
                
                # Initialize with default value
                first_visit["Aclaraciones"] = "-"
                
                parts = []
                if ":" in block[aclaraciones_idx]:
                    # Extract value after colon in same line
                    value = block[aclaraciones_idx].split(":", 1)[1].strip()
                    if value:
                        parts.append(value)
                
                # Collect text from following lines
                while j < len(block) and not block[j].startswith("Próxima visita") and not block[j].startswith("Solicitud para la autorización de pruebas"):
                    parts.append(block[j])
                    j += 1
                
                first_visit["Aclaraciones"] = "\n\n".join(parts) if parts else "-"
                    
            except ValueError:
                first_visit["Aclaraciones"] = "-"
                
            # Merge with existing self.first_medical_visit row if it exists, otherwise assign.
            new_first = pd.DataFrame([first_visit])
            if hasattr(self, "first_medical_visit") and not self.first_medical_visit.empty:
                # Update the existing row with non-empty values from new_first.
                for col in new_first.columns:
                    new_val = new_first.iloc[0][col]
                    if new_val not in ["", None]:
                        self.first_medical_visit.at[0, col] = new_val
            else:
                self.first_medical_visit = new_first

        # Ensure self.next_medical_visits exists as a list.
        if not hasattr(self, "next_medical_visits") or not isinstance(self.next_medical_visits, list):
            self.next_medical_visits = []

        ################ Process subsequent blocks (extra visits) ################
        for idx, block in enumerate(blocks[1:]):
            visit = {}
            # Extract treatment.
            treat_idx = None
            treat_header = None
            for header in ["Tratamiento y evolución. Exploraciones complementarias", "Evolución"]:
                try:
                    t_idx = block.index(header)
                    if treat_idx is None or t_idx < treat_idx:
                        treat_idx = t_idx
                        treat_header = header
                except ValueError:
                    continue
            if treat_idx is not None:
                j = treat_idx + 1
                treatment_parts = []
                while j < len(block) and block[j] != "Estado actual y exploración física":
                    treatment_parts.append(block[j])
                    j += 1
                treatment_text = "\n\n".join(treatment_parts)
                visit["Tratamiento y evolución. Exploraciones complementarias"] = treatment_text
                # For "Evolución", use the previous element for "Fecha de consulta extra"
                if treat_header == "Evolución" and treat_idx > 0:
                    visit["Fecha de consulta extra"] = block[treat_idx - 1][:10]
                else:
                    visit["Fecha de consulta extra"] = ""
            else:
                visit["Tratamiento y evolución. Exploraciones complementarias"] = "-"
                visit["Fecha de consulta extra"] = ""
            # Extract state information.
            try:
                state_idx = block.index("Estado actual y exploración física")
                j = state_idx + 1
                state_parts = []
                while j < len(block) and block[j] != "Relación de causalidad":
                    state_parts.append(block[j])
                    j += 1
                state_text = " ".join(state_parts)
                
                # Initialize with default values
                visit["HISTORIA ACTUAL"] = "-"
                visit["EXPLORACION FISICA"] = "-" 
                visit["Pruebas complementarias"] = "-"
                
                # Try to extract each section
                sections = state_text.split("HISTORIA ACTUAL:")
                if len(sections) > 1:
                    remaining = sections[1].strip()
                    historia = remaining.split("EXPLORACION FISICA:")[0].strip()
                    visit["HISTORIA ACTUAL"] = historia
                    
                    sections = remaining.split("EXPLORACION FISICA:")
                    if len(sections) > 1:
                        remaining = sections[1].strip()
                        exploracion = remaining.split("Pruebas complementarias:")[0].strip()
                        visit["EXPLORACION FISICA"] = exploracion
                        
                        sections = remaining.split("Pruebas complementarias:")
                        if len(sections) > 1:
                            pruebas = sections[1].strip()
                            visit["Pruebas complementarias"] = pruebas
            except ValueError:
                visit["Estado actual y exploración física"] = "-"
                
                
            # Extract lesiones temporales information
            try:
                lesiones_idx = block.index("Lesiones temporales")
                j = lesiones_idx + 1
                
                # Initialize with default values
                visit["Intervenciones quirúrgicas"] = "-"
                visit["Patrimonial. Daño emergente (se indemniza su importe)"] = "-"
                
                while j < len(block) and block[j] != "Secuelas. Básico":
                    text = block[j]
                    
                    # Check for "Intervenciones quirúrgicas"
                    if text.startswith("Intervenciones quirúrgicas"):
                        parts = []
                        if ":" in text:
                            # Extract value after colon in same line
                            value = text.split(":", 1)[1].strip()
                            if value:
                                parts.append(value)
                        # Also collect text from following lines
                        j += 1
                        while j < len(block) and not block[j].startswith("Patrimonial. Daño emergente"):
                            parts.append(block[j])
                            j += 1
                        visit["Intervenciones quirúrgicas"] = " ".join(parts) if parts else "-"
                        continue # Skip j increment since we already moved forward
                        
                    # Check for "Patrimonial. Daño emergente" 
                    elif text.startswith("Patrimonial. Daño emergente"):
                        parts = []
                        if ":" in text:
                            # Extract value after colon in same line
                            value = text.split(":", 1)[1].strip()
                            if value:
                                parts.append(value)
                        # Also collect text from following lines
                        j += 1
                        while j < len(block) and block[j] != "Secuelas. Básico":
                            parts.append(block[j])
                            j += 1
                        visit["Patrimonial. Daño emergente (se indemniza su importe)"] = "\n\n".join(parts) if parts else "-"
                        continue # Skip j increment since we already moved forward
                    
                    j += 1
                    
            except ValueError:
                visit["Intervenciones quirúrgicas"] = "-"
                visit["Patrimonial. Daño emergente (se indemniza su importe)"] = "-"
                
            # Extract Valoración Total Secuelas information
            try:
                j = 0
                while j < len(block):
                    if "Valoración Total Secuelas" in block[j]:
                        text = block[j]
                        if ":" in text:
                            value = text.split(":", 1)[1].strip()
                            visit["Valoración Total Secuelas"] = value if value else "-"
                            
                            # Check next line for "Motivos variación"
                            if j+1 < len(block) and "Motivos variación" in block[j+1]:
                                motivos_text = block[j+1]
                                if ":" in motivos_text:
                                    motivos_value = motivos_text.split(":", 1)[1].strip()
                                    visit["Motivos variación"] = motivos_value if motivos_value else "-"
                        else:
                            visit["Valoración Total Secuelas"] = "-"
                        break
                    j += 1
                    
                if "Valoración Total Secuelas" not in visit:
                    visit["Valoración Total Secuelas"] = "-"
                
                if "Motivos variación" not in visit:
                    visit["Motivos variación"] = "-"
                    
            except Exception:
                visit["Valoración Total Secuelas"] = "-"
                visit["Motivos variación"] = "-"
            
            # Extract aclaraciones information
            try:
                aclaraciones_idx = block.index("Aclaraciones:")
                j = aclaraciones_idx + 1
                
                # Initialize with default value
                visit["Aclaraciones"] = "-"
                
                parts = []
                if ":" in block[aclaraciones_idx]:
                    # Extract value after colon in same line
                    value = block[aclaraciones_idx].split(":", 1)[1].strip()
                    if value:
                        parts.append(value)
                
                # Collect text from following lines
                while j < len(block) and not block[j].startswith("Próxima visita") and not block[j].startswith("Solicitud para la autorización de pruebas"):
                    parts.append(block[j])
                    j += 1
                
                visit["Aclaraciones"] = "\n\n".join(parts) if parts else "-"
                    
            except ValueError:
                visit["Aclaraciones"] = "-"
                
            new_visit = pd.DataFrame([visit])
            # Merge with existing DataFrame for this extra visit index if it exists.
            if len(self.next_medical_visits) > idx and not self.next_medical_visits[idx].empty:
                for col in new_visit.columns:
                    new_val = new_visit.iloc[0][col]
                    if new_val not in ["", "-", None]:
                        self.next_medical_visits[idx].at[0, col] = new_val
            else:
                self.next_medical_visits.append(new_visit)
